from fastapi import FastAPI, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv
import os
import json
from typing import Dict, Any
import logging
import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
logger.info("Environment variables loaded")

# Initialize FastAPI app
app = FastAPI(title="NoteMap AI Backend")
logger.info("FastAPI application initialized")

# Configure CORS with all necessary origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:3001", 
        "http://localhost:3002",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:3001",
        "http://127.0.0.1:3002"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=3600,
)
logger.info("CORS middleware configured with all development origins")

# Configure Gemini AI with safety settings
GOOGLE_API_KEY = "AIzaSyBfG56Iqoi0tv_yB4LZ3rIfpDXtcc8MsOk"
if not GOOGLE_API_KEY:
    logger.error("GOOGLE_API_KEY not found in environment variables")
    raise RuntimeError("GOOGLE_API_KEY not found in environment variables")
else:
    logger.info("Google API key configured successfully")

try:
    # Configure the Gemini API
    genai.configure(api_key=GOOGLE_API_KEY)
    logger.info("Gemini API configured successfully")
    
    # List available models
    logger.info("Fetching available Gemini models...")
    models = genai.list_models()
    for m in models:
        if "gemini" in m.name.lower():
            logger.info(f"Found Gemini model: {m.name}")
    
    # Configure the model with specific parameters
    generation_config = {
        "temperature": 0.3,
        "top_p": 0.8,
        "top_k": 40,
        "max_output_tokens": 2048,
    }

    safety_settings = [
        {
            "category": "HARM_CATEGORY_HARASSMENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_HATE_SPEECH",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
    ]

    model = genai.GenerativeModel(
        model_name="gemini-1.5-pro",  # Using the latest stable version
        generation_config=generation_config,
        safety_settings=safety_settings
    )
    logger.info("Gemini model initialized successfully")
    
except Exception as e:
    logger.error(f"Failed to initialize Gemini API: {str(e)}", exc_info=True)
    raise RuntimeError(f"Failed to initialize Gemini API: {str(e)}")

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        logger.info("Starting PDF text extraction")
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            text += page_text
            logger.debug(f"Extracted text from page {page_num}")
        logger.info(f"Successfully extracted {len(text)} characters from {doc.page_count} pages")
        return text
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        logger.info("Starting DOCX text extraction")
        from io import BytesIO
        doc = Document(BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        logger.info(f"Successfully extracted {len(text)} characters from DOCX with {len(doc.paragraphs)} paragraphs")
        return text
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

async def generate_mind_map(text: str) -> Dict[str, Any]:
    """Generate mind map structure using Gemini AI."""
    if not text.strip():
        logger.warning("Empty text provided to generate_mind_map")
        raise HTTPException(status_code=400, detail="Empty text provided")
        
    logger.info(f"Starting mind map generation for text of length: {len(text)}")
    
    prompt = """You are an expert at creating hierarchical mind maps from educational content. Your task is to analyze the given text and create a clear, well-structured mind map in JSON format.

Instructions:
1. Identify the main topic/concept from the text
2. Break down the content into major subtopics
3. For each subtopic, identify key points and supporting details
4. Organize the information in a hierarchical structure
5. Ensure the structure is balanced and logical
6. Keep node names concise but descriptive
7. Include important relationships between concepts

The response MUST be a valid JSON object with exactly this structure:
{
    "root": "Main Topic",
    "branches": [
        {
            "name": "Subtopic 1",
            "branches": [
                {"name": "Key Point 1.1"},
                {"name": "Key Point 1.2"}
            ]
        }
    ]
}

Text to analyze:
"""
    
    try:
        logger.info("Preparing Gemini API request")
        print(prompt)
        logger.debug(f"Using prompt template with text sample: {text[:200]}...")
        
        response = model.generate_content(
            contents=prompt + "\n" + text
        )
        
        logger.info("Received response from Gemini API")
        
        # Check if response was blocked
        if response.prompt_feedback.block_reason:
            logger.error(f"Content blocked by Gemini API: {response.prompt_feedback.block_reason}")
            raise HTTPException(
                status_code=400, 
                detail=f"Content blocked: {response.prompt_feedback.block_reason}"
            )

        # Extract JSON from the response
        response_text = response.text
        logger.debug(f"Raw response from Gemini API: {response_text[:200]}...")
        
        # Find the JSON part
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0]
            logger.debug("Found JSON response in code block")
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0]
            logger.debug("Found JSON response in generic code block")
        else:
            json_str = response_text
            logger.debug("Using raw response as JSON")
        
        # Parse and validate JSON
        try:
            mind_map = json.loads(json_str)
            if not isinstance(mind_map, dict) or "root" not in mind_map or "branches" not in mind_map:
                logger.error("Invalid mind map structure received")
                logger.debug(f"Received structure: {json_str}")
                raise ValueError("Invalid mind map structure")
            
            logger.info(f"Successfully generated mind map with root topic: {mind_map['root']}")
            logger.debug(f"Mind map has {len(mind_map['branches'])} main branches")
            return mind_map
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {str(e)}", exc_info=True)
            logger.debug(f"Failed JSON string: {json_str}")
            raise HTTPException(status_code=500, detail=f"Failed to parse AI response as JSON: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error generating mind map: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating mind map: {str(e)}")

@app.post("/api/upload")
async def upload_file(file: UploadFile):
    """
    Upload and process a file to generate a mind map.
    Accepts PDF, DOCX, or TXT files.
    """
    if not file:
        logger.warning("No file provided in upload request")
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    logger.info(f"Processing uploaded file: {file.filename} (type: {file.content_type})")
    
    try:
        # Read file content
        content = await file.read()
        logger.info(f"Successfully read {len(content)} bytes from uploaded file")
        
        # Extract text based on file type
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(content)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(content)
        elif file.filename.endswith('.txt'):
            text = content.decode('utf-8')
            logger.info(f"Successfully extracted {len(text)} characters from TXT file")
        else:
            logger.warning(f"Unsupported file type: {file.filename}")
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT files.")
        
        # Generate mind map
        logger.info("Starting mind map generation")
        mind_map = await generate_mind_map(text)
        logger.info("Mind map generation completed successfully")
        
        return {"mind_map": mind_map}
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    status = {
        "status": "healthy",
        "api_key_configured": bool(GOOGLE_API_KEY),
        "timestamp": datetime.datetime.now().isoformat()
    }
    logger.info(f"Health check request - Status: {status}")
    return status 